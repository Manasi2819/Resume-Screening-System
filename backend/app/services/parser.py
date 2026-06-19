"""
Multi-format document parser — supports PDF, DOCX, image (OCR), and plain text.

Design:
  extract_text(file_bytes, filename)  ← unified dispatcher (main entry point)
      │
      ├─ .pdf   → extract_text_from_pdf   (pdfplumber)
      ├─ .docx  → extract_text_from_docx  (python-docx)
      ├─ .doc   → raises ValueError with user-friendly message
      ├─ .png/.jpg/.jpeg/.webp → extract_text_from_image (pytesseract OCR)
      └─ .txt   → extract_text_from_txt   (UTF-8 decode)

After text extraction the pipeline is unchanged:
  chunk_by_sections → extract_metadata → parse_resume
"""
import io
import logging
import os
import re
from pathlib import Path
from typing import Optional

import pdfplumber
import spacy

logger = logging.getLogger(__name__)

# ── Optional imports — fail gracefully with clear messages ───────────────────
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore

try:
    from PIL import Image
    import pytesseract
    _TESSERACT_AVAILABLE = True
except ImportError:
    _TESSERACT_AVAILABLE = False
    pytesseract = None  # type: ignore
    Image = None  # type: ignore

# ── Configure Tesseract path from .env (Windows local dev) ───────────────────
_settings = None
try:
    from app.core.config import settings as _settings
    if _settings.TESSERACT_CMD and _TESSERACT_AVAILABLE:
        if os.path.exists(_settings.TESSERACT_CMD):
            pytesseract.pytesseract.tesseract_cmd = _settings.TESSERACT_CMD
except Exception:
    pass  # settings not available at import time in some test contexts

# ── Load spaCy model once at module level ────────────────────────────────────
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    raise RuntimeError(
        "spaCy model not found. Run: python -m spacy download en_core_web_sm"
    )

# ── Accepted extensions ───────────────────────────────────────────────────────
RESUME_ALLOWED_EXTS = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".webp"}
JD_ALLOWED_EXTS     = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".webp", ".txt"}

# ── Section header keywords ───────────────────────────────────────────────────
SECTION_HEADERS: dict[str, list[str]] = {
    "summary": ["summary", "objective", "profile", "about me", "career objective"],
    "skills": [
        "skills", "technical skills", "core competencies", "technologies",
        "tools", "tech stack", "areas of expertise",
        # Portal-specific (Naukri / LinkedIn screenshots)
        "key skills", "keyskills", "may also know", "top skills",
    ],
    "experience": [
        "experience", "work experience", "employment", "professional experience",
        "work history", "career history",
        # Portal-specific
        "current", "previous",
    ],
    "education": [
        "education", "academic background", "qualifications", "academics",
        # Portal-specific
        "educational details",
    ],
    "projects": ["projects", "personal projects", "key projects", "portfolio"],
    "certifications": ["certifications", "certificates", "awards", "achievements"],
    "location": [
        # Portal-specific — captures preferred location lines
        "pref. locations", "preferred locations", "pref locations",
        "pret locations", "pret ocatons",  # common OCR corruptions
    ],
}

# ── Common tech skills for metadata extraction ────────────────────────────────
SKILL_KEYWORDS = [
    "python", "java", "javascript", "typescript", "react", "node.js", "nodejs",
    "sql", "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
    "machine learning", "deep learning", "nlp", "computer vision", "llm",
    "fastapi", "django", "flask", "spring boot", "express",
    "docker", "kubernetes", "aws", "gcp", "azure", "terraform",
    "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "langchain",
    "git", "ci/cd", "rest api", "graphql", "kafka", "spark",
    "c++", "c#", "go", "rust", "scala", "r",
]


# ════════════════════════════════════════════════════════════════════════════════
# Text sanitization helpers
# ════════════════════════════════════════════════════════════════════════════════

def heal_split_headers(text: str) -> str:
    """Heal PDF layout splitting issues where first letters are on a separate line (drop caps)."""
    lines = text.split("\n")
    if len(lines) < 2:
        return text

    i = 0
    while i < len(lines) - 1:
        line1 = lines[i].strip()
        line2 = lines[i + 1].strip()
        words1 = line1.split()
        words2 = line2.split()

        if (len(words1) >= 2 and
                all(len(w) == 1 and w.isupper() for w in words1) and
                len(words1) == len(words2)):
            healed_words = [w1 + w2 for w1, w2 in zip(words1, words2)]
            lines[i] = " ".join(healed_words)
            lines.pop(i + 1)
            continue
        i += 1

    return "\n".join(lines)


def sanitize_text(text: str) -> str:
    """
    Remove characters that PostgreSQL cannot store in text columns:
    - NUL bytes (0x00) — PostgreSQL string literals cannot contain them
    - Other non-printable control characters (except \\t, \\n, \\r)
    """
    text = text.replace("\x00", "")
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    return text


# ════════════════════════════════════════════════════════════════════════════════
# Format-specific extractors
# ════════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber. Returns plain text string. Falls back to OCR if scanned."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Could not read PDF: {e}")

    text = sanitize_text(text)

    # Check if we should fallback to OCR (scanned PDF)
    if len(text.strip()) < 100:
        if not _TESSERACT_AVAILABLE:
            raise ValueError(
                "The PDF appears to be a scanned image or has no machine-readable text layer. "
                "OCR (pytesseract/Pillow) is not available to extract text. Please install pytesseract/Pillow or upload a text-based PDF."
            )

        ocr_text = ""
        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    # Render the page as image
                    im = page.to_image(resolution=150)
                    pil_img = im.original
                    if pil_img.mode not in ("RGB", "L"):
                        pil_img = pil_img.convert("RGB")
                    page_ocr = pytesseract.image_to_string(pil_img, lang="eng")
                    ocr_text += page_ocr + "\n"
            text = ocr_text
        except Exception as e:
            raise ValueError(
                f"The PDF has no text layer, and OCR extraction failed. Error: {e}"
            )

    text = sanitize_text(text)
    return heal_split_headers(text.strip())


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file using python-docx."""
    if DocxDocument is None:
        raise ValueError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Table cells (skills tables, education grids, etc.)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        text = "\n".join(parts)
    except Exception as e:
        raise ValueError(f"Could not read DOCX file: {e}")

    return sanitize_text(text.strip())


def _preprocess_image_for_ocr(image: "Image.Image") -> "Image.Image":  # type: ignore[name-defined]
    """
    Preprocess an image to improve Tesseract OCR accuracy.

    Steps:
      1. Upscale to at least 1800 px wide (Tesseract needs ~300 DPI equiv.)
      2. Convert to grayscale
      3. Apply adaptive thresholding via ImageFilter to sharpen text

    This dramatically improves accuracy on portal screenshots, which are
    typically low-resolution, have coloured backgrounds, and use small fonts.
    """
    from PIL import ImageFilter, ImageOps

    # 1. Upscale if too small — Tesseract struggles below ~150 DPI
    min_width = 1800
    if image.width < min_width:
        scale = min_width / image.width
        new_size = (int(image.width * scale), int(image.height * scale))
        image = image.resize(new_size, Image.LANCZOS)

    # 2. Grayscale
    image = ImageOps.grayscale(image)

    # 3. Sharpen to make text edges crisper
    image = image.filter(ImageFilter.SHARPEN)

    return image


def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Extract text from an image (PNG/JPG/WEBP) using Tesseract OCR.

    Applies preprocessing (upscale → grayscale → sharpen) before OCR to
    improve accuracy on portal screenshots and low-resolution images.
    Uses PSM 6 (uniform block of text) for structured portal-style layouts.
    """
    if not _TESSERACT_AVAILABLE:
        raise ValueError(
            "pytesseract / Pillow is not installed. "
            "Run: pip install pytesseract Pillow"
        )
    try:
        image = Image.open(io.BytesIO(file_bytes))
        # Convert to RGB if needed (handles RGBA PNGs, etc.)
        if image.mode not in ("RGB", "L"):
            image = image.convert("RGB")

        # Preprocess for better OCR on screenshots / low-res images
        preprocessed = _preprocess_image_for_ocr(image)

        # PSM 6: assume a single uniform block of text (good for portal cards)
        ocr_config = r"--oem 3 --psm 6"
        text = pytesseract.image_to_string(preprocessed, lang="eng", config=ocr_config)
    except Exception as e:
        raise ValueError(
            f"Could not OCR image. Make sure Tesseract is installed and "
            f"TESSERACT_CMD is set correctly in .env if on Windows. Error: {e}"
        )

    return sanitize_text(text.strip())


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode a plain-text file as UTF-8 (fallback to latin-1)."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="replace")
    return sanitize_text(text.strip())


# ════════════════════════════════════════════════════════════════════════════════
# Unified dispatcher
# ════════════════════════════════════════════════════════════════════════════════

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Route to the correct extractor based on file extension.

    Accepted:
      .pdf                        → pdfplumber
      .docx                       → python-docx
      .doc                        → rejected (user-friendly error)
      .png / .jpg / .jpeg / .webp → Vision LLM with Tesseract OCR fallback
      .txt                        → UTF-8 decode (JD only — routes.py enforces this)

    Raises ValueError with a descriptive message on any failure.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)

    if ext == ".docx":
        return extract_text_from_docx(file_bytes)

    if ext == ".doc":
        raise ValueError(
            f"'{filename}' — .doc files are not supported. "
            "Please open the file in Microsoft Word (or LibreOffice), "
            "choose 'Save As', select '.docx' format, and re-upload."
        )

    if ext in {".png", ".jpg", ".jpeg", ".webp"}:
        vision_enabled = _settings.USE_VISION_FOR_IMAGES if _settings else False
        if vision_enabled:
            try:
                from app.services import vision_parser
                jd_text = vision_parser.extract_jd_with_vision(
                    image_bytes=file_bytes,
                    groq_api_key=_settings.GROQ_API_KEY,
                    vision_model=_settings.GROQ_VISION_MODEL,
                )
                if jd_text is not None:
                    return jd_text
            except Exception as e:
                logger.warning(
                    f"[extract_text] JD Vision path raised exception — using OCR fallback. Detail: {e}"
                )
        return extract_text_from_image(file_bytes)

    if ext == ".txt":
        return extract_text_from_txt(file_bytes)

    raise ValueError(
        f"Unsupported file type '{ext}' for file '{filename}'. "
        "Accepted formats — Resumes: PDF, DOCX, PNG, JPG, JPEG, WEBP. "
        "Job Descriptions: PDF, DOCX, TXT, PNG, JPG, JPEG, WEBP."
    )


# ════════════════════════════════════════════════════════════════════════════════
# Section-aware chunking
# ════════════════════════════════════════════════════════════════════════════════

def detect_section(line: str) -> Optional[str]:
    """Return section label if this line is a section header, else None."""
    cleaned = line.lower().strip().rstrip(":").strip()
    for section, keywords in SECTION_HEADERS.items():
        if cleaned in keywords:
            return section
        cleaned_no_special = re.sub(r"[^a-z\s]", "", cleaned).strip()
        if cleaned_no_special in keywords:
            return section
    return None


def chunk_by_sections(text: str) -> list[dict]:
    """
    Split resume text into labelled section chunks.
    Returns: [{section: str, text: str}, ...]
    """
    lines = text.split("\n")
    chunks: list[dict] = []
    current_section = "summary"
    current_lines: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        section = detect_section(stripped)
        if section and len(stripped) < 60:
            if current_lines:
                chunk_text = "\n".join(current_lines).strip()
                if len(chunk_text) > 30:
                    chunks.append({"section": current_section, "text": chunk_text})
            current_section = section
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        chunk_text = "\n".join(current_lines).strip()
        if len(chunk_text) > 30:
            chunks.append({"section": current_section, "text": chunk_text})

    return chunks


# ════════════════════════════════════════════════════════════════════════════════
# Metadata extraction (name, email, phone, skills, years experience)
# ════════════════════════════════════════════════════════════════════════════════

def _strip_ocr_garbage(text: str) -> str:
    """
    Strip leading non-ASCII garbage characters that Tesseract produces from
    portal screenshots (e.g. checkbox icons rendered as '\ufffd', bullets, etc.).
    Keeps only printable ASCII plus common Unicode letters.
    """
    # Remove leading non-letter characters (OCR artefacts like '\ufffd', bullets)
    cleaned = re.sub(r"^[^\w]+", "", text.strip())
    # Also collapse internal runs of non-printable chars
    cleaned = re.sub(r"[\x00-\x1f\x7f-\x9f\ufffd]+", " ", cleaned).strip()
    return cleaned


def heuristic_extract_name(text: str) -> str:
    """Extract candidate name from first 10 lines of text using formatting filters."""
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    forbidden_keywords = {
        "resume", "cv", "curriculum", "vitae", "summary", "profile", "contact",
        "email", "phone", "mobile", "address", "linkedin", "github", "page",
        "portfolio", "experience", "education", "skills", "projects", "objective",
        # Portal-specific forbidden words
        "current", "previous", "fresher", "pune", "mumbai", "bangalore",
        "bengaluru", "hyderabad", "delhi", "chennai", "similar", "profiles",
        "comment", "save", "view", "number", "candidate", "lacs", "location",
        "locations", "preferred", "pref",
    }

    for line in lines[:10]:
        # Strip leading OCR garbage (icons, checkboxes, bullets from portal UIs)
        cleaned = _strip_ocr_garbage(line)
        if not (3 <= len(cleaned) <= 50):
            continue
        if "@" in cleaned or ".com" in cleaned or ".in" in cleaned:
            continue
        digits_count = sum(c.isdigit() for c in cleaned)
        if digits_count > 3:
            continue
        words = re.findall(r"\b\w+\b", cleaned.lower())
        if any(w in forbidden_keywords for w in words):
            continue
        if len(words) > 5:  # names rarely exceed 5 words
            continue
        job_titles = {"engineer", "developer", "analyst", "student", "architect",
                      "lead", "manager", "intern", "associate", "consultant"}
        if any(w in job_titles for w in words):
            continue
        if "http" in cleaned.lower() or "www." in cleaned.lower():
            continue
        if re.match(r"^[A-Z][a-zA-Z\s\.\-]+$", cleaned):
            return cleaned
        if cleaned.isupper() and re.match(r"^[A-Z\s\.\-]+$", cleaned):
            return cleaned

    # Fallback to spaCy NER
    doc = nlp(text[:3000])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            val = ent.text.strip()
            val = re.sub(r"[\|;\+0-9\ufffd]", "", val).strip()
            if len(val) >= 3 and not any(
                w in forbidden_keywords for w in re.findall(r"\b\w+\b", val.lower())
            ):
                return val

    # Ultimate fallback: first non-garbage line of text
    for line in lines[:5]:
        first_line = _strip_ocr_garbage(line)
        if len(first_line) > 2:
            return first_line

    return "Unknown"


# ── Known Indian cities for location extraction ──────────────────────────────
_CITY_NAMES = [
    "pune", "mumbai", "bangalore", "bengaluru", "hyderabad", "chennai",
    "delhi", "noida", "gurgaon", "gurugram", "kolkata", "ahmedabad",
    "jaipur", "indore", "bhopal", "nagpur", "surat", "lucknow",
    "coimbatore", "kochi", "thiruvananthapuram", "chandigarh",
]


def _extract_years_experience(text: str) -> int:
    """
    Extract years of experience supporting multiple formats:
      - "3 years of experience" / "3+ years"
      - Portal format: "3y 6m" / "1y" / "0y 6m"
    """
    # Standard resume format
    matches = re.findall(
        r"(\d+)\+?\s*years?\s*(?:of\s+)?(?:experience|exp)?",
        text,
        re.IGNORECASE,
    )
    if matches:
        return max(int(x) for x in matches)

    # Portal card format: "Xy Xm" (e.g. "3y 6m", "0y 6m", "1y")
    portal_match = re.search(r"(\d+)\s*y\s*(?:\d+\s*m)?", text, re.IGNORECASE)
    if portal_match:
        years = int(portal_match.group(1))
        # If 0y, check months — if months >= 6 treat as ~0.5 yr to avoid hard gate
        if years == 0:
            months_match = re.search(r"0\s*y\s*(\d+)\s*m", text, re.IGNORECASE)
            if months_match and int(months_match.group(1)) >= 6:
                return 1  # treat 6+ months as 1 year to pass the experience gate
        return years

    return 0


def _extract_location(text: str) -> Optional[str]:
    """Extract candidate's preferred/current location from resume or portal text."""
    text_lower = text.lower()
    found = [city.title() for city in _CITY_NAMES if city in text_lower]
    return ", ".join(dict.fromkeys(found)) if found else None  # deduplicate, preserve order


def extract_metadata(text: str) -> dict:
    """
    Extract structured metadata from raw resume text using spaCy + regex.
    Returns: {name, email, phone, years_experience, skills, location}
    """
    doc = nlp(text[:3000])  # noqa: F841 — kept for potential future NER use

    name = heuristic_extract_name(text)

    emails = re.findall(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}", text)
    email = emails[0] if emails else None

    phones = re.findall(r"[\+]?[\d\s\-\(\)]{10,16}", text)
    phone = phones[0].strip() if phones else None

    years_exp = _extract_years_experience(text)

    text_lower = text.lower()
    found_skills = [s for s in SKILL_KEYWORDS if s in text_lower]

    location = _extract_location(text)

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "years_experience": years_exp,
        "skills": found_skills,
        "location": location,
    }


# ════════════════════════════════════════════════════════════════════════════════
# Main entry point
# ════════════════════════════════════════════════════════════════════════════════

def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Full parse pipeline for a resume file:
      raw bytes + filename → {raw_text, chunks, metadata}

    Routing logic:
      Image files (.png/.jpg/.jpeg/.webp)
        ├─ USE_VISION_FOR_IMAGES=true → Vision LLM (Groq Llama 4 Scout)
        │     ├─ Success → rich structured metadata + clean raw_text
        │     └─ Any failure (rate limit, network, bad JSON, size >4MB)
        │           → silently falls back to improved OCR path
        └─ USE_VISION_FOR_IMAGES=false → OCR only
      PDF / DOCX / TXT → existing parsers (unchanged)

    The `parsed_by` key in metadata records which path was used:
      "vision_llm" | "ocr" | "pdf" | "docx" | "txt"
    """
    ext = Path(filename).suffix.lower()

    # ── Vision LLM path for image files ────────────────────────────────
    if ext in {".png", ".jpg", ".jpeg", ".webp"}:
        vision_enabled = getattr(_settings, "USE_VISION_FOR_IMAGES", True)
        if vision_enabled:
            try:
                from app.services import vision_parser
                vision_data = vision_parser.extract_profile_with_vision(
                    image_bytes=file_bytes,
                    groq_api_key=_settings.GROQ_API_KEY,
                    vision_model=_settings.GROQ_VISION_MODEL,
                )
                if vision_data is not None:
                    metadata = vision_parser.vision_data_to_metadata(vision_data)
                    # Use Vision LLM's reconstructed text; fall back to OCR if empty
                    raw_text = sanitize_text((vision_data.get("raw_text") or "").strip())
                    if len(raw_text) < 50:
                        logger.info("[parse_resume] Vision raw_text too short — supplementing with OCR text")
                        raw_text = extract_text_from_image(file_bytes)
                    chunks = chunk_by_sections(raw_text)
                    if not chunks:
                        chunks = [{"section": "experience", "text": raw_text[:3000]}]
                    return {"raw_text": raw_text, "chunks": chunks, "metadata": metadata}
                # Vision returned None — fall through to OCR below
                logger.info("[parse_resume] Vision returned None — using OCR fallback")
            except Exception as e:
                logger.warning(f"[parse_resume] Vision path raised exception — using OCR fallback. Detail: {e}")

    # ── Standard path: PDF / DOCX / TXT / image-OCR fallback ──────────────
    raw_text = extract_text(file_bytes, filename)
    chunks = chunk_by_sections(raw_text)
    metadata = extract_metadata(raw_text)

    # Stamp which parser was used
    _ext_to_parser = {".pdf": "pdf", ".docx": "docx", ".txt": "txt"}
    metadata["parsed_by"] = _ext_to_parser.get(ext, "ocr")

    if not chunks:
        chunks = [{"section": "experience", "text": raw_text[:3000]}]

    return {
        "raw_text": raw_text,
        "chunks": chunks,
        "metadata": metadata,
    }
